#-*- coding: utf-8 -*-
##############################################################################
#
# Copyright (c) 2017 Shoobx, Inc.
# All Rights Reserved.
#
# This software is subject to the provisions of the Zope Public License,
# Version 2.1 (ZPL).  A copy of the ZPL should accompany this distribution.
# THIS SOFTWARE IS PROVIDED "AS IS" AND ANY AND ALL EXPRESS OR IMPLIED
# WARRANTIES ARE DISCLAIMED, INCLUDING, BUT NOT LIMITED TO, THE IMPLIED
# WARRANTIES OF TITLE, MERCHANTABILITY, AGAINST INFRINGEMENT, AND FITNESS
# FOR A PARTICULAR PURPOSE.
#
##############################################################################
"""Flowable Element Processing
"""
import docx
import lazy
import lxml
import re
import zope.interface
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.style import WD_STYLE
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import RGBColor
from docx.text.run import Font, Run
from z3c.rml import directive, occurence
from z3c.rml import flowable as rml_flowable
from z3c.rml import template as rml_template

# from z3c.rml flowable.py file
from z3c.rml import attr, directive, interfaces, platypus
try:
    import reportlab.graphics.barcode
except ImportError:
    # barcode package has not been installed
    import types
    import reportlab.graphics
    reportlab.graphics.barcode = types.ModuleType('barcode')
    reportlab.graphics.barcode.createBarcodeDrawing = None


def pygments2xpre(s, language="python"):
    "Return markup suitable for XPreformatted"
    try:
        from pygments import highlight
        from pygments.formatters import HtmlFormatter
    except ImportError:
        return s

    from pygments.lexers import get_lexer_by_name

    l = get_lexer_by_name(language)

    h = HtmlFormatter()
    # XXX: Does not work in Python 2, since pygments creates non-unicode
    # outpur snippets.
    #from io import StringIO
    from six import StringIO
    out = StringIO()
    highlight(s,l,h,out)
    styles = [(cls, style.split(';')[0].split(':')[1].strip())
                for cls, (style, ttype, level) in h.class2style.items()
                if cls and style and style.startswith('color:')]
    from reportlab.lib.pygments2xpre import _2xpre
    return _2xpre(out.getvalue(),styles)


class Flowable(directive.RMLDirective):
    klass=None
    attrMapping = None

    @property
    def container(self):
        # Goes up the tree to find the content container in order to
        # append new paragraph
        parent = self.parent
        while not IContentContainer.providedBy(parent):
            parent = parent.parent
        return parent.container

    def process(self):
        args = dict(self.getAttributeValues(attrMapping=self.attrMapping))


class Spacer(Flowable):
    signature = rml_flowable.ISpacer
    klass = reportlab.platypus.Spacer
    attrMapping = {'length': 'height'}


class Illustration(Flowable):
    signature = rml_flowable.IIllustration
    klass = platypus.Illustration

    def process(self):
        args = dict(self.getAttributeValues())


class BarCodeFlowable(Flowable):
    signature = rml_flowable.IBarCodeFlowable
    klass = staticmethod(reportlab.graphics.barcode.createBarcodeDrawing)
    attrMapping = {'code': 'codeName'}


class SubParagraphDirective(directive.RMLDirective):

    @lazy.lazy
    def paragraph(self):
        para = self.parent
        while not IParagraph.providedBy(para):
            para = para.parent
        return para


class IComplexSubParagraphDirective(interfaces.IRMLDirectiveSignature):
    """A sub-paragraph directive that can contian further elements."""

class ComplexSubParagraphDirective(SubParagraphDirective):
    """A sub-paragraph directive that can contian further elements."""
    signature = IComplexSubParagraphDirective
    factories = {}

    def __init__(self, element, parent):
        super(SubParagraphDirective, self).__init__(element, parent)
        self.originalStyles = {}

    def setStyle(self, name, value):
        self.originalStyles[name] = getattr(self.paragraph, name)
        setattr(self.paragraph, name, value)

    def unsetStyles(self):
        for name, value in self.originalStyles.items():
            setattr(self.paragraph, name, value)

    def preProcess(self):
        pass

    def postProcess(self):
        self.unsetStyles()

    def process(self):
        self.preProcess()
        if self.element.text:
            self.paragraph.addRun(self.element.text)
        self.processSubDirectives()
        self.postProcess()

        if self.element.tail:
            self.paragraph.addRun(self.element.tail)


class IItalic(IComplexSubParagraphDirective):
    """Renders the text inside as italic."""

class Italic(ComplexSubParagraphDirective):
    signature = IItalic

    def preProcess(self):
        self.setStyle('italic', True)


class IBold(IComplexSubParagraphDirective):
    """Renders the text inside as bold."""

class Bold(ComplexSubParagraphDirective):
    signature = IBold

    def preProcess(self):
        self.setStyle('bold', True)


class IUnderline(IComplexSubParagraphDirective):
    """Renders the text inside as underline."""

class Underline(ComplexSubParagraphDirective):
    signature = IUnderline

    def preProcess(self):
        self.setStyle('underline', True)


class IStrike(IComplexSubParagraphDirective):
    """Renders the text inside as strike."""

class Strike(ComplexSubParagraphDirective):
    signature = IStrike

    def preProcess(self):
        self.setStyle('strike', True)


class IFont(IComplexSubParagraphDirective):
    """Renders the text inside as strike."""

    face = attr.String(
        title=u'Font face',
        description=u'The name of the font for the cell.',
        required=False)

    size = attr.Measurement(
        title=u'Font Size',
        description=u'The font size for the text of the cell.',
        required=False)

    color = attr.Color(
        title=u'Font Color',
        description=u'The color in which the text will appear.',
        required=False)

class Font(ComplexSubParagraphDirective):
    signature = IFont

    def preProcess(self):
        data = self.getAttributeValues(
            attrMapping={'face': 'fontName',
                         'size': 'fontSize',
                         'color': 'fontColor'})
        for name, value in data:
            self.setStyle(name, value)


class ISuper(IComplexSubParagraphDirective):
    """Renders the text inside as super.

    Note that a cusotm font size and position is not supported.
    """

class Super(ComplexSubParagraphDirective):
    signature = ISuper

    def preProcess(self):
        self.setStyle('superscript', True)


class ISub(IComplexSubParagraphDirective):
    """Renders the text inside as sub.

    Note that a cusotm font size and position is not supported.
    """

class Sub(ComplexSubParagraphDirective):
    signature = ISub

    def preProcess(self):
        self.setStyle('subscript', True)


class IBreak(interfaces.IRMLDirectiveSignature):
    """Adds a break in the paragraph.
    """

class Break(SubParagraphDirective):
    signature = IBreak

    def process(self):
        run = self.paragraph.addRun()
        run.add_break(WD_BREAK.LINE)

        if self.element.tail:
            self.paragraph.addRun(self.element.tail)


class IPageNumber(interfaces.IRMLDirectiveSignature):
    """Adds a break in the paragraph.
    """

class PageNumber(SubParagraphDirective):
    signature = IPageNumber

    def process(self):
        run = self.paragraph.addRun()
        # XXX: TO BE DONE!!!

        if self.element.tail:
            self.paragraph.addRun(self.element.tail)


class IAnchor(IComplexSubParagraphDirective):
    """Adds an anchor link into the paragraph."""

    backcolor = attr.Color(
        title=u'Background Color',
        description=u'The background color of the link area.',
        required=False)

    color = attr.Color(
        title=u'Font Color',
        description=u'The color in which the text will appear.',
        required=False)

    url = attr.Text(
        title=u'URL',
        description=u'The URL to link to.',
        required=False)

    fontName = attr.String(
        title=u'Font face',
        description=u'The name of the font for the cell.',
        required=False)

    fontSize = attr.Measurement(
        title=u'Font Size',
        description=u'The font size for the text of the cell.',
        required=False)

    name = attr.Text(
        title=u'Name',
        description=u'The name of the link.',
        required=False)


class Anchor(ComplexSubParagraphDirective):
    signature = IAnchor

    def process(self):
        attrs = dict(self.getAttributeValues())
        # This gets access to the document.xml.rels file and gets
        # a new relation id value
        part = self.paragraph.docxParagraph.part
        r_id = part.relate_to(
            attrs['url'],
            docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
            is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

        # Create a w:r element
        run = docx.oxml.shared.OxmlElement('w:r')

        # Create a new w:rPr element
        rPr = docx.oxml.shared.OxmlElement('w:rPr')

        # Sets the color
        color = docx.oxml.shared.OxmlElement('w:color')
        color.set(docx.oxml.shared.qn('w:val'), '0000FF')
        rPr.append(color)

        # Underlines the text
        underline = docx.oxml.shared.OxmlElement('w:u')
        underline.set(docx.oxml.shared.qn('w:val'), 'single')
        rPr.append(underline)

        # Join all the xml elements together add add the required
        # text to the w:r element
        run.append(rPr)
        run.text = self.element.text
        hyperlink.append(run)
        self.paragraph.docxParagraph._p.append(hyperlink)


ComplexSubParagraphDirective.factories = {
    'i': Italic,
    'b': Bold,
    'u': Underline,
    'strike': Strike,
    'strong': Bold,
    'font': Font,
    'super': Super,
    'sub': Sub,
    'a': Anchor,
    'br': Break,
    'pageNumber': PageNumber,
    # Unsupported tags:
    # 'greek': Greek,
}

IComplexSubParagraphDirective.setTaggedValue(
    'directives',
    (occurence.ZeroOrMore('i', IItalic),
     occurence.ZeroOrMore('b', IBold),
     occurence.ZeroOrMore('u', IUnderline),
     occurence.ZeroOrMore('strike', IStrike),
     occurence.ZeroOrMore('strong', IBold),
     occurence.ZeroOrMore('font', IFont),
     occurence.ZeroOrMore('super', ISuper),
     occurence.ZeroOrMore('sub', ISub),
     occurence.ZeroOrMore('a', IAnchor),
     occurence.ZeroOrMore('br', IBreak),
    )
)

class IParagraph(zope.interface.Interface):
    pass

@zope.interface.implementer(IParagraph)
class Paragraph(Flowable):
    signature = rml_flowable.IParagraph
    defaultStyle = 'Normal'
    factories = {
        'i': Italic,
        'b': Bold,
        'u': Underline,
        'strike': Strike,
        'strong': Bold,
        'font': Font,
        'super': Super,
        'sub': Sub,
        'a': Anchor,
        'br': Break,
        'pageNumber': PageNumber,
        # Unsupported tags:
        # 'greek': Greek,
    }

    italic = None
    bold = None
    underline = None
    strike = None
    fontName = None
    fontSize  = None
    fontColor = None
    superscript = None
    subscript = None

    overrideStyle = None

    def _cleanText(self, text):
        if not text:
            text = ''
        text = re.sub('\n\s+', ' ', text)
        text = re.sub('\s\s\s+', '', text)
        text = re.sub('\t', '', text)
        return text

    def addRun(self, text=None):
        if text is not None:
            text = self._cleanText(text)
        run = self.docxParagraph.add_run(text)
        run.font.italic = self.italic
        run.font.bold = self.bold
        run.font.underline = self.underline
        run.font.strike = self.strike
        run.font.name = self.fontName
        run.font.size  = self.fontSize
        if self.fontColor is not None:
            run.font.color.rgb = RGBColor(
                *[int(c*255) for c in self.fontColor.rgb()])
        if self.superscript is not None:
            run.font._element.rPr.superscript = self.superscript
        if self.subscript is not None:
            run.font.subscript = self.subscript
        return run

    def process(self):
        # Styles within li tags are given to paras as attributes
        # This retrieves and applies the given style
        style = self.element.attrib.get('style', self.defaultStyle)

        # Append new paragraph.
        self.docxParagraph = self.container.add_paragraph(style=style)

        if self.element.text:
            self.addRun(self.element.text)

        self.processSubDirectives()


class Preformatted(Paragraph):
    signature = rml_flowable.IPreformatted
    klass = reportlab.platypus.Preformatted


class XPreformatted(Paragraph):
    signature = rml_flowable.IXPreformatted
    klass = reportlab.platypus.XPreformatted


class Heading1(Paragraph):
    signature = rml_flowable.IHeading1
    defaultStyle = "Heading1"


class Heading2(Paragraph):
    signature = rml_flowable.IHeading2
    defaultStyle = "Heading2"
    overrideStyle = None


class Heading3(Paragraph):
    signature = rml_flowable.IHeading3
    defaultStyle = "Heading3"


class Heading4(Paragraph):
    signature = rml_flowable.IHeading4
    defaultStyle = "Heading4"


class Heading5(Paragraph):
    signature = rml_flowable.IHeading5
    defaultStyle = "Heading5"


class Heading6(Paragraph):
    signature = rml_flowable.IHeading6
    defaultStyle = "Heading6"


class Link(Flowable):
    signature = rml_flowable.ILink
    attrMapping = {'destination': 'destinationname',
                   'boxStrokeWidth': 'thickness',
                   'boxStrokeDashArray': 'dashArray',
                   'boxStrokeColor': 'color'}

    defaultStyle = "Normal"

    def process(self):
        # Takes care of case where li object does not have <para> tag
        children = self.element.getchildren()
        if len(children) == 0 or children[0].tag != 'para':
            newPara = lxml.etree.Element('para')
            newPara.text = self.element.text
            self.element.text = None
            for subElement in tuple(self.element): newPara.append(subElement)
            self.element.append(newPara)
        flow = Flow(self.element, self.parent)
        flow.process()


class HorizontalRow(Flowable):
    signature = rml_flowable.IHorizontalRow
    klass = reportlab.platypus.flowables.HRFlowable
    attrMapping = {'align': 'hAlign'}

    def _handleText(self, element, hr):
        run = hr.add_run(element)

    def process(self):
        # Implement other alignment styles? self.element.attrib has values
        hr = self.parent.container.add_paragraph()
        hr_format = hr.paragraph_format
        hr_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        docx_bar = u'───────────────────────────────────────────────────────────'
        self._handleText(docx_bar, hr)
        return hr


class Title(directive.RMLDirective):
    signature = rml_flowable.ITitle
    defaultStyle = "Title"

    def _cleanText(self, text):
        if not text:
            text = ''
        text = re.sub('\n\s+', ' ', text)
        text = re.sub('\s\s\s+', '', text)
        text = re.sub('\t', '', text)
        return text

    def _handleText(self, element, title):
        if element.text is not None and element.text.strip() != '':
            run = title.add_run(self._cleanText(element.text.lstrip()))

    def process(self):
        style = self.element.attrib.get('style', self.defaultStyle)
        # Note: Using add_heading instead of add_paragraph. Does this still inherit?
        title = self.parent.container.add_heading(level = 0)
        self._handleText(self.element, title)
        return title


class Flow(directive.RMLDirective):
    factories = {
        # Generic Flowables
        'spacer': Spacer,
        'illustration': Illustration,
        'barCodeFlowable': BarCodeFlowable,
        'pre': Preformatted,
        #'xpre': XPreformatted,
        # Paragraph-Like Flowables
        'para': Paragraph,
        # Headers
        'h1': Heading1,
        'h2': Heading2,
        'h3': Heading3,
        'h4': Heading4, 
        'h5': Heading5,
        'h6': Heading6,
        'title': Title,
        'hr':HorizontalRow,
        'link': Link
    }

    def __init__(self, *args, **kw):
        super(Flow, self).__init__(*args, **kw)

    def process(self):
        self.processSubDirectives()
